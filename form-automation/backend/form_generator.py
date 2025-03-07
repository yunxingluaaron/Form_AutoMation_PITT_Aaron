# form_generator.py
from fpdf import FPDF
from docx import Document
import io
import os
from datetime import datetime

class FormGenerator:
    """Base class for form generation"""
    
    def __init__(self, output_dir='./generated_forms'):
        self.output_dir = output_dir
        os.makedirs(output_dir, exist_ok=True)
    
    def generate_filename(self, prefix, patient_name):
        """Generate a unique filename based on patient name and timestamp"""
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        safe_name = "".join(c if c.isalnum() else "_" for c in patient_name)
        return f"{prefix}_{safe_name}_{timestamp}"


class IBHSFormGenerator(FormGenerator):
    """Generator for IBHS Form"""
    
    def generate_pdf(self, form_data):
        """
        Generate IBHS form as PDF
        
        Args:
            form_data (dict): Mapped form data
            
        Returns:
            str: Path to generated PDF file
        """
        pdf = FPDF()
        pdf.add_page()
        
        # Set up the PDF
        pdf.set_font("Helvetica", "B", 16)
        pdf.cell(0, 10, "WRITTEN ORDER FOR IBHS", 0, 1, "C")
        pdf.ln(10)
        
        # Add form fields
        pdf.set_font("Helvetica", "B", 12)
        pdf.cell(0, 10, "Child Information", 0, 1)
        
        pdf.set_font("Helvetica", "", 10)
        pdf.cell(40, 10, "Child's Name:", 0, 0)
        pdf.cell(0, 10, form_data.get("child_name", ""), 0, 1)
        
        pdf.cell(40, 10, "Date of Birth:", 0, 0)
        pdf.cell(0, 10, form_data.get("dob", ""), 0, 1)
        
        pdf.cell(40, 10, "Parent/Guardian:", 0, 0)
        pdf.cell(0, 10, form_data.get("parent_guardian", ""), 0, 1)
        
        pdf.ln(5)
        
        # Diagnoses
        pdf.set_font("Helvetica", "B", 12)
        pdf.cell(0, 10, "CURRENT BEHAVIORAL HEALTH DIAGNOSES", 0, 1)
        pdf.set_font("Helvetica", "", 10)
        
        # Handle multi-line text for diagnoses
        diagnoses = form_data.get("current_diagnoses", "").split("\n")
        for diagnosis in diagnoses:
            pdf.multi_cell(0, 10, diagnosis)
        
        pdf.ln(5)
        
        # Goals
        pdf.set_font("Helvetica", "B", 12)
        pdf.cell(0, 10, "MEASURABLE GOALS AND OBJECTIVES", 0, 1)
        pdf.set_font("Helvetica", "", 10)
        
        goals = form_data.get("measurable_goals", "").split("\n")
        for goal in goals:
            pdf.multi_cell(0, 10, goal)
        
        pdf.ln(5)
        
        # Clinical information
        pdf.set_font("Helvetica", "B", 12)
        pdf.cell(0, 10, "CLINICAL INFORMATION", 0, 1)
        pdf.set_font("Helvetica", "", 10)
        
        clinical_info = form_data.get("clinical_information", "").split("\n")
        for info in clinical_info:
            pdf.multi_cell(0, 10, info)
        
        # Signature section
        pdf.ln(20)
        pdf.cell(0, 10, "Prescriber Signature: _______________________", 0, 1)
        pdf.cell(0, 10, f"Date: {datetime.now().strftime('%m/%d/%Y')}", 0, 1)
        
        # Save the PDF
        filename = self.generate_filename("IBHS_Form", form_data.get("child_name", "Patient"))
        file_path = os.path.join(self.output_dir, f"{filename}.pdf")
        pdf.output(file_path)
        
        return file_path
    
    def generate_docx(self, form_data):
        """
        Generate IBHS form as DOCX
        
        Args:
            form_data (dict): Mapped form data
            
        Returns:
            str: Path to generated DOCX file
        """
        doc = Document()
        
        # Add title
        doc.add_heading("WRITTEN ORDER FOR IBHS", 0)
        
        # Add child information
        doc.add_heading("Child Information", level=1)
        doc.add_paragraph(f"Child's Name: {form_data.get('child_name', '')}")
        doc.add_paragraph(f"Date of Birth: {form_data.get('dob', '')}")
        doc.add_paragraph(f"Parent/Guardian: {form_data.get('parent_guardian', '')}")
        
        # Add diagnoses
        doc.add_heading("CURRENT BEHAVIORAL HEALTH DIAGNOSES", level=1)
        doc.add_paragraph(form_data.get("current_diagnoses", ""))
        
        # Add goals
        doc.add_heading("MEASURABLE GOALS AND OBJECTIVES", level=1)
        doc.add_paragraph(form_data.get("measurable_goals", ""))
        
        # Add clinical information
        doc.add_heading("CLINICAL INFORMATION", level=1)
        doc.add_paragraph(form_data.get("clinical_information", ""))
        
        # Add signature section
        doc.add_paragraph("")
        doc.add_paragraph("")
        doc.add_paragraph(f"Prescriber Signature: _______________________")
        doc.add_paragraph(f"Date: {datetime.now().strftime('%m/%d/%Y')}")
        
        # Save the document
        filename = self.generate_filename("IBHS_Form", form_data.get("child_name", "Patient"))
        file_path = os.path.join(self.output_dir, f"{filename}.docx")
        doc.save(file_path)
        
        return file_path


class CommunityCareFormGenerator(FormGenerator):
    """Generator for Community Care IBHS Form"""
    
    def generate_pdf(self, form_data):
        """
        Generate Community Care form as PDF
        
        Args:
            form_data (dict): Mapped form data
            
        Returns:
            str: Path to generated PDF file
        """
        pdf = FPDF()
        pdf.add_page()
        
        # Set up the PDF
        pdf.set_font("Helvetica", "B", 16)
        pdf.cell(0, 10, "COMMUNITY CARE IBHS WRITTEN ORDER LETTER", 0, 1, "C")
        pdf.ln(10)
        
        # Add current date
        pdf.set_font("Helvetica", "", 10)
        pdf.cell(0, 10, f"Date: {datetime.now().strftime('%B %d, %Y')}", 0, 1)
        pdf.ln(5)
        
        # Add recipient information
        pdf.set_font("Helvetica", "B", 12)
        pdf.cell(0, 10, "Recipient Information", 0, 1)
        
        pdf.set_font("Helvetica", "", 10)
        pdf.cell(40, 10, "Name:", 0, 0)
        pdf.cell(0, 10, form_data.get("recipient_name", ""), 0, 1)
        
        pdf.cell(40, 10, "DOB:", 0, 0)
        pdf.cell(0, 10, form_data.get("dob", ""), 0, 1)
        
        pdf.cell(40, 10, "Age:", 0, 0)
        pdf.cell(0, 10, form_data.get("age", ""), 0, 1)
        
        pdf.ln(5)
        
        # Diagnoses
        pdf.set_font("Helvetica", "B", 12)
        pdf.cell(0, 10, "Current Diagnoses", 0, 1)
        pdf.set_font("Helvetica", "", 10)
        
        diagnoses = form_data.get("diagnoses", "").split("\n")
        for diagnosis in diagnoses:
            pdf.multi_cell(0, 10, diagnosis)
        
        pdf.ln(5)
        
        # Symptoms
        pdf.set_font("Helvetica", "B", 12)
        pdf.cell(0, 10, "Clinical Presentation/Symptoms", 0, 1)
        pdf.set_font("Helvetica", "", 10)
        
        symptoms = form_data.get("symptoms", "").split("\n")
        for symptom in symptoms:
            pdf.multi_cell(0, 10, symptom)
        
        pdf.ln(5)
        
        # Treatment recommendations
        pdf.set_font("Helvetica", "B", 12)
        pdf.cell(0, 10, "Treatment Recommendations", 0, 1)
        pdf.set_font("Helvetica", "", 10)
        
        recommendations = form_data.get("treatment_recommendations", "").split("\n")
        for rec in recommendations:
            pdf.multi_cell(0, 10, rec)
        
        pdf.ln(5)
        
        # Medical necessity
        pdf.set_font("Helvetica", "B", 12)
        pdf.cell(0, 10, "Medical Necessity", 0, 1)
        pdf.set_font("Helvetica", "", 10)
        
        necessity = form_data.get("medical_necessity", "").split("\n")
        for line in necessity:
            pdf.multi_cell(0, 10, line)
        
        # Signature section
        pdf.ln(20)
        pdf.cell(0, 10, "Provider Signature: _______________________", 0, 1)
        pdf.cell(0, 10, "Provider Name (Printed): _______________________", 0, 1)
        pdf.cell(0, 10, f"Date: {datetime.now().strftime('%m/%d/%Y')}", 0, 1)
        
        # Save the PDF
        filename = self.generate_filename("Community_Care_Form", form_data.get("recipient_name", "Patient"))
        file_path = os.path.join(self.output_dir, f"{filename}.pdf")
        pdf.output(file_path)
        
        return file_path
    
    def generate_docx(self, form_data):
        """
        Generate Community Care form as DOCX
        
        Args:
            form_data (dict): Mapped form data
            
        Returns:
            str: Path to generated DOCX file
        """
        doc = Document()
        
        # Add title
        doc.add_heading("COMMUNITY CARE IBHS WRITTEN ORDER LETTER", 0)
        
        # Add date
        doc.add_paragraph(f"Date: {datetime.now().strftime('%B %d, %Y')}")
        
        # Add recipient information
        doc.add_heading("Recipient Information", level=1)
        doc.add_paragraph(f"Name: {form_data.get('recipient_name', '')}")
        doc.add_paragraph(f"DOB: {form_data.get('dob', '')}")
        doc.add_paragraph(f"Age: {form_data.get('age', '')}")
        
        # Add diagnoses
        doc.add_heading("Current Diagnoses", level=1)
        doc.add_paragraph(form_data.get("diagnoses", ""))
        
        # Add symptoms
        doc.add_heading("Clinical Presentation/Symptoms", level=1)
        doc.add_paragraph(form_data.get("symptoms", ""))